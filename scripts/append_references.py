"""Append a `References` section to the Capturing-Science-Talk doc.

Pulls the two references already cited in the .docx comments and adds a small
set of additional appropriate references. Saves in place. The original file is
preserved by the sibling .backup.docx written manually before running this.

Run from the project root:
    python scripts/append_references.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = PROJECT_ROOT / "Captureing Science Talk in Musuem_5.12.26+DI.docx"

REFERENCES = [
    # Alphabetical by first author. Plain hyphens for page ranges (no em/en
    # dashes, per the project's style preference).
    (
        "Gilkerson, J., Richards, J. A., Warren, S. F., Montgomery, J. K., "
        "Greenwood, C. R., Kimbrough Oller, D., Hansen, J. H. L., & Paul, T. D. "
        "(2017). Mapping the early language environment using all-day recordings "
        "and automated analysis. American Journal of Speech-Language Pathology, "
        "26(2), 248-265. https://doi.org/10.1044/2016_AJSLP-15-0169"
    ),
    (
        "Grattafiori, A., et al. (2024). The Llama 3 herd of models "
        "(arXiv:2407.21783) [Preprint]. arXiv. https://arxiv.org/abs/2407.21783"
    ),
    (
        "Haden, C. A., Jant, E. A., Hoffman, P. C., Marcus, M., Geddes, J. R., "
        "& Gaskins, S. (2014). Supporting family conversations and children's "
        "STEM learning in a children's museum. Early Childhood Research "
        "Quarterly, 29(3), 333-344. https://doi.org/10.1016/j.ecresq.2014.04.004"
    ),
    (
        "Jameel, M. I., & Dungen, J. (2015). Low-power wireless advertising "
        "software library for distributed M2M and contextual IoT. IEEE, 597-602."
    ),
    (
        "Karpukhin, V., Oguz, B., Min, S., Lewis, P., Wu, L., Edunov, S., "
        "Chen, D., & Yih, W. (2020). Dense passage retrieval for open-domain "
        "question answering. In Proceedings of the 2020 Conference on Empirical "
        "Methods in Natural Language Processing (EMNLP) (pp. 6769-6781). "
        "Association for Computational Linguistics. "
        "https://doi.org/10.18653/v1/2020.emnlp-main.550"
    ),
    (
        "Radford, A., Kim, J. W., Xu, T., Brockman, G., McLeavey, C., "
        "Sutskever, I., & Zaremba, W. (2022). Robust speech recognition via "
        "large-scale weak supervision (arXiv:2212.04356) [Preprint]. arXiv. "
        "https://arxiv.org/abs/2212.04356"
    ),
    (
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings "
        "using Siamese BERT-networks. In Proceedings of the 2019 Conference on "
        "Empirical Methods in Natural Language Processing and the 9th "
        "International Joint Conference on Natural Language Processing "
        "(EMNLP-IJCNLP) (pp. 3982-3992). Association for Computational "
        "Linguistics. https://doi.org/10.18653/v1/D19-1410"
    ),
]


def append_references(doc_path: Path = DOC_PATH) -> Path:
    doc = Document(doc_path)

    heading = doc.add_heading("References", level=2)
    for run in heading.runs:
        run.font.size = Pt(12)

    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.size = Pt(11)

    doc.save(doc_path)
    return doc_path


if __name__ == "__main__":
    out = append_references()
    print(f"Updated {out}")
