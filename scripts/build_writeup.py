"""Generate the science-talk-detection write-up as a .docx file.

Run from the project root:
    python scripts/build_writeup.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "Science_Talk_Detection_Writeup.docx"


TITLE = "Detecting Science Talk in Parent-Child Conversations at a Museum"

SECTIONS = [
    (
        "Overview",
        [
            (
                "Museums are some of the richest informal learning spaces a young child "
                "ever gets to experience. When a parent and child wander through an "
                "exhibit, the questions they ask, the predictions they trade, and the "
                "small explanations they offer each other are exactly the moments "
                "researchers and educators care about. Capturing them at scale is "
                "genuinely hard, though: hours of audio sit on a recording device, and "
                "most of what is said is small talk, logistics, or redirection. The "
                "science is in there, but it is buried. This write-up describes the "
                "system being built to surface those moments automatically, end to end, "
                "from raw audio to a researcher-ready list of flagged utterances."
            ),
        ],
    ),
    (
        "System Architecture",
        [
            (
                "The run-time pipeline has four stages: audio capture, speech "
                "recognition, text normalization, and the science-talk detector itself. "
                "Each stage hands a clean artifact to the next, so any one component "
                "can be swapped without rewiring the others."
            ),
            (
                "1. Audio capture. The parent carries a small recording device during "
                "the visit. It records continuously and stores a single audio file per "
                "session; nothing is processed on the device."
            ),
            (
                "2. Automatic speech recognition. The uploaded audio is transcribed "
                "with Whisper-large-v3, OpenAI's open-source ASR model. We chose it "
                "because it holds up well in noisy, real-world recordings (background "
                "chatter, exhibit sounds, the occasional sneeze) and handles child "
                "speech noticeably better than most alternatives we evaluated. The "
                "output is a time-stamped transcript with rough speaker turns. We treat "
                "this as a best-effort text layer rather than ground truth, and the "
                "downstream components are designed to tolerate the errors Whisper "
                "tends to make."
            ),
            (
                "3. Text normalization. Before any modeling, the transcript goes "
                "through a light cleanup pass: Unicode normalization, whitespace "
                "collapse, removal of obvious filler, and segmentation into "
                "utterance-level rows keyed to a timestamp and speaker. This is the "
                "same normalization used on the training corpus, which keeps training "
                "and inference text in the same shape."
            ),
            (
                "4. Science-talk detector. Each utterance is passed through a "
                "fine-tuned bi-encoder, producing a dense embedding. That embedding is "
                "compared (cosine similarity) against an anchor bank of curated "
                "science-talk utterances, each tagged with one of five subtypes: "
                "observation, prediction, causal reasoning, evidence, or content. If "
                "the top-match similarity clears a tuned threshold, the utterance is "
                "flagged and inherits the subtype of its nearest anchor. Borderline "
                "cases are sent through a secondary LLM gate that asks, in plain "
                "language, whether the utterance counts as science talk. The gate is "
                "the same one used to score hard negatives during training, so its "
                "behavior is already calibrated."
            ),
            (
                "The final output is a structured table per visit: timestamp, speaker, "
                "verbatim utterance, predicted subtype, and a confidence score. "
                "Researchers can sort and filter the table, jump back to the audio at "
                "any timestamp, and optionally confirm or correct the label, which "
                "feeds future training rounds."
            ),
        ],
    ),
    (
        "How the Detector Was Trained",
        [
            (
                "The detector is a retrieval-style model trained on a small but "
                "carefully constructed corpus of early-childhood science utterances. "
                "The corpus is built in four stages."
            ),
            (
                "First, ingestion of a labeled seed corpus: roughly 200 expert-coded "
                "utterances drawn from prior research on early-childhood science talk, "
                "each carrying the lexical cues that signaled science to the original "
                "coders."
            ),
            (
                "Second, sub-type labeling. Science talk is not one thing, so we split "
                "it into the five subtypes above using a three-stage classifier: "
                "deterministic rules over the explicit cues, a keyword scan over a "
                "curated science vocabulary, and an LLM zero-shot fallback for any "
                "utterance the first two stages cannot handle."
            ),
            (
                "Third, negative mining. The hardest part of this problem is the "
                "negatives: a parent saying \"I wonder where your shoes are\" looks "
                "suspiciously like science. We mine three kinds of negatives: clean "
                "transcript utterances with no science cues; LLM-generated hard "
                "negatives that mirror the syntax of a positive but live in a "
                "non-science topic; and seed-term negatives where a science word like "
                "\"predict\" or \"what if\" is used in a clearly non-scientific way. "
                "The negative pool ends up more than three times the size of the "
                "positive pool, balanced across subtypes."
            ),
            (
                "Fourth, conversational paraphrase augmentation. The seed utterances "
                "came from classroom settings; museum talk is casual, side-by-side, "
                "parent-to-child. We use an LLM to rewrite each positive into the kind "
                "of short, informal phrasing a parent would actually use with their "
                "own child, while preserving the science cues. This stretches the "
                "corpus and pulls it toward the deployment setting. The augmented "
                "corpus is then used to fine-tune the bi-encoder, and the same LLM "
                "gate that scores hard negatives during training serves as the "
                "borderline-case gate at inference."
            ),
        ],
    ),
    (
        "Why This Approach",
        [
            (
                "Two constraints shaped the design. The seed corpus is small, so "
                "training a large model from scratch is not an option; we have to "
                "stretch a few hundred utterances into something that generalizes. And "
                "the cost of false positives is high, because every wrongly flagged "
                "utterance is researcher time wasted. The four-stage training pipeline "
                "addresses both pressures, and the run-time architecture keeps the "
                "system simple to operate: a parent uploads a file and gets back a "
                "reviewable list of probable science moments, each tagged by subtype."
            ),
        ],
    ),
]


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _tighten_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def build_document(output_path: Path = OUTPUT_PATH) -> Path:
    doc = Document()
    _set_default_font(doc)
    _tighten_margins(doc)

    title = doc.add_heading(TITLE, level=1)
    for run in title.runs:
        run.font.size = Pt(15)

    for heading, paragraphs in SECTIONS:
        h = doc.add_heading(heading, level=2)
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after = Pt(2)
        for run in h.runs:
            run.font.size = Pt(12)
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Wrote {path}")
